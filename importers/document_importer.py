"""
Document importer for Seshat.

Imports text from Word documents, PDFs, and other document formats.
"""

from typing import Dict, List, Any, Optional, Generator
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path


@dataclass
class ImportedDocument:
    """Represents an imported document."""
    text: str
    filename: str
    file_type: str
    title: Optional[str]
    author: Optional[str]
    created_date: Optional[datetime]
    modified_date: Optional[datetime]
    page_count: Optional[int]
    word_count: int
    metadata: Dict[str, Any]

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            "text": self.text,
            "filename": self.filename,
            "file_type": self.file_type,
            "title": self.title,
            "author": self.author,
            "created_date": self.created_date.isoformat() if self.created_date else None,
            "modified_date": self.modified_date.isoformat() if self.modified_date else None,
            "page_count": self.page_count,
            "word_count": self.word_count,
            "metadata": self.metadata,
        }


class DocumentImporter:
    """
    Import text from various document formats.

    Supports PDF, DOCX, TXT, and other common formats.
    """

    SUPPORTED_EXTENSIONS = {
        ".txt": "text",
        ".md": "markdown",
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "doc",
        ".rtf": "rtf",
        ".odt": "odt",
    }

    def __init__(self):
        """Initialize document importer."""
        pass

    def import_file(self, file_path: str) -> Optional[ImportedDocument]:
        """
        Import a document file.

        Args:
            file_path: Path to document

        Returns:
            ImportedDocument or None if failed
        """
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            return None

        file_type = self.SUPPORTED_EXTENSIONS[ext]

        if file_type in ("text", "markdown"):
            return self._import_text(path)
        elif file_type == "pdf":
            return self._import_pdf(path)
        elif file_type == "docx":
            return self._import_docx(path)
        elif file_type == "doc":
            return self._import_doc(path)
        elif file_type == "rtf":
            return self._import_rtf(path)
        elif file_type == "odt":
            return self._import_odt(path)

        return None

    def import_directory(
        self,
        directory: str,
        recursive: bool = True,
        limit: Optional[int] = None,
    ) -> Generator[ImportedDocument, None, None]:
        """
        Import all documents from a directory.

        Args:
            directory: Directory path
            recursive: Search subdirectories
            limit: Maximum files to import

        Yields:
            ImportedDocument for each file
        """
        path = Path(directory)
        count = 0

        for ext in self.SUPPORTED_EXTENSIONS:
            if limit and count >= limit:
                break

            pattern = f"**/*{ext}" if recursive else f"*{ext}"

            for file_path in path.glob(pattern):
                if limit and count >= limit:
                    break

                doc = self.import_file(str(file_path))
                if doc:
                    yield doc
                    count += 1

    def _import_text(self, path: Path) -> Optional[ImportedDocument]:
        """Import plain text or markdown file."""
        try:
            text = path.read_text(encoding="utf-8")
            stat = path.stat()

            return ImportedDocument(
                text=text,
                filename=path.name,
                file_type="text",
                title=path.stem,
                author=None,
                created_date=datetime.fromtimestamp(stat.st_ctime),
                modified_date=datetime.fromtimestamp(stat.st_mtime),
                page_count=None,
                word_count=len(text.split()),
                metadata={"encoding": "utf-8"},
            )

        except Exception:
            try:
                text = path.read_text(encoding="latin-1")
                return ImportedDocument(
                    text=text,
                    filename=path.name,
                    file_type="text",
                    title=path.stem,
                    author=None,
                    created_date=None,
                    modified_date=None,
                    page_count=None,
                    word_count=len(text.split()),
                    metadata={"encoding": "latin-1"},
                )
            except Exception:
                return None

    def _import_pdf(self, path: Path) -> Optional[ImportedDocument]:
        """Import PDF file."""
        try:
            import pdfplumber

            text_parts = []
            page_count = 0
            metadata = {}

            with pdfplumber.open(path) as pdf:
                page_count = len(pdf.pages)

                if pdf.metadata:
                    metadata = dict(pdf.metadata)

                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

            text = "\n\n".join(text_parts)

            return ImportedDocument(
                text=text,
                filename=path.name,
                file_type="pdf",
                title=metadata.get("Title", path.stem),
                author=metadata.get("Author"),
                created_date=self._parse_pdf_date(metadata.get("CreationDate")),
                modified_date=self._parse_pdf_date(metadata.get("ModDate")),
                page_count=page_count,
                word_count=len(text.split()),
                metadata=metadata,
            )

        except ImportError:
            try:
                import PyPDF2

                text_parts = []

                with open(path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    page_count = len(reader.pages)
                    metadata = reader.metadata or {}

                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)

                text = "\n\n".join(text_parts)

                return ImportedDocument(
                    text=text,
                    filename=path.name,
                    file_type="pdf",
                    title=metadata.get("/Title", path.stem),
                    author=metadata.get("/Author"),
                    created_date=None,
                    modified_date=None,
                    page_count=page_count,
                    word_count=len(text.split()),
                    metadata=dict(metadata) if metadata else {},
                )

            except Exception:
                return None

        except Exception:
            return None

    def _import_docx(self, path: Path) -> Optional[ImportedDocument]:
        """Import DOCX file."""
        try:
            from docx import Document

            doc = Document(path)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            text = "\n\n".join(text_parts)

            core_props = doc.core_properties
            metadata = {}

            if hasattr(core_props, "author"):
                metadata["author"] = core_props.author
            if hasattr(core_props, "title"):
                metadata["title"] = core_props.title

            return ImportedDocument(
                text=text,
                filename=path.name,
                file_type="docx",
                title=getattr(core_props, "title", None) or path.stem,
                author=getattr(core_props, "author", None),
                created_date=getattr(core_props, "created", None),
                modified_date=getattr(core_props, "modified", None),
                page_count=None,
                word_count=len(text.split()),
                metadata=metadata,
            )

        except Exception:
            return None

    def _import_doc(self, path: Path) -> Optional[ImportedDocument]:
        """Import legacy DOC file (requires antiword or similar)."""
        import subprocess

        try:
            result = subprocess.run(
                ["antiword", str(path)],
                capture_output=True,
                text=True,
                timeout=30,
            )

            if result.returncode == 0:
                text = result.stdout

                return ImportedDocument(
                    text=text,
                    filename=path.name,
                    file_type="doc",
                    title=path.stem,
                    author=None,
                    created_date=None,
                    modified_date=None,
                    page_count=None,
                    word_count=len(text.split()),
                    metadata={"converter": "antiword"},
                )

        except Exception:
            pass

        return None

    def _import_rtf(self, path: Path) -> Optional[ImportedDocument]:
        """Import RTF file."""
        try:
            from striprtf.striprtf import rtf_to_text

            rtf_content = path.read_text(encoding="utf-8", errors="ignore")
            text = rtf_to_text(rtf_content)

            return ImportedDocument(
                text=text,
                filename=path.name,
                file_type="rtf",
                title=path.stem,
                author=None,
                created_date=None,
                modified_date=None,
                page_count=None,
                word_count=len(text.split()),
                metadata={},
            )

        except Exception:
            return None

    def _import_odt(self, path: Path) -> Optional[ImportedDocument]:
        """Import ODT (OpenDocument Text) file."""
        try:
            from odf import text as odf_text
            from odf.opendocument import load

            doc = load(path)
            paragraphs = doc.getElementsByType(odf_text.P)

            text_parts = []
            for para in paragraphs:
                para_text = ""
                for node in para.childNodes:
                    if node.nodeType == node.TEXT_NODE:
                        para_text += node.data
                    elif hasattr(node, "childNodes"):
                        for child in node.childNodes:
                            if child.nodeType == child.TEXT_NODE:
                                para_text += child.data

                if para_text.strip():
                    text_parts.append(para_text)

            text = "\n\n".join(text_parts)

            return ImportedDocument(
                text=text,
                filename=path.name,
                file_type="odt",
                title=path.stem,
                author=None,
                created_date=None,
                modified_date=None,
                page_count=None,
                word_count=len(text.split()),
                metadata={},
            )

        except Exception:
            return None

    def _parse_pdf_date(self, date_str: Optional[str]) -> Optional[datetime]:
        """Parse PDF date string."""
        if not date_str:
            return None

        try:
            if date_str.startswith("D:"):
                date_str = date_str[2:]

            date_str = date_str[:14]

            return datetime.strptime(date_str, "%Y%m%d%H%M%S")

        except Exception:
            return None
